import csv
import sys
from dateutil.parser import parse
import win32com.client
import time
from flask import (
    Flask,
    request,
    jsonify,
    send_file,
    send_from_directory,
    render_template,
    make_response,
    Blueprint,    
)
from flask_cors import CORS
import pandas as pd
import os
from docx import Document
from io import BytesIO
import tempfile
import zipfile
import re
from datetime import datetime
import uuid
import traceback
import shutil
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import threading
import logging
import queue
import logging
import traceback

logging.basicConfig(level=logging.DEBUG)

document_bp = Blueprint("document", __name__)
CORS(document_bp)

EMAIL_RATE_LIMIT = 2  # Number of emails to send per time period
EMAIL_TIME_PERIOD = 10  # Time period in seconds
# Create directories if they don't exist
# Global state
generation_progress = {}
generated_files = {}
email_progress = {}
email_failures = {}

# Helper functions (format_date, is_date_field, parse_date, replace_placeholder_in_docx) remain the same

from flask import render_template

@document_bp.route("/")
def document_home():
    return render_template("document_creation.html")

def format_date(date, format_str):
    if format_str == "DD/MM/YYYY":
        return date.strftime("%d/%m/%Y")
    elif format_str == "MM/DD/YYYY":
        return date.strftime("%m/%d/%Y")
    elif format_str == "DD MMM YYYY":
        return date.strftime("%d %b %Y")
    elif format_str == "DD MMMM YYYY":
        return date.strftime("%d %B %Y")
    else:
        return date.strftime("%Y-%m-%d")


def is_date_field(key, value):
    """Check if a field is likely a date field"""
    date_keywords = [
        "date",
        "Date",
        "DATE",
        "dob",
        "DOB",
        "birth",
        "Birth",
        "BIRTH",
        "expiry",
        "Expiry",
        "EXPIRY",
        "issued",
        "Issued",
        "ISSUED",
    ]

    if any(keyword in key for keyword in date_keywords):
        return True

    date_patterns = [
        r"\d{4}-\d{2}-\d{2}",
        r"\d{2}/\d{2}/\d{4}",
        r"\d{2}-\d{2}-\d{4}",
        r"\d{1,2} \w{3,9} \d{4}",
    ]

    if isinstance(value, str):
        for pattern in date_patterns:
            if re.fullmatch(pattern, value.strip()):
                return True

    return False


def parse_date(date_str):
    """Try to parse a date string into a datetime object"""
    try:
        return parse(date_str)
    except (ValueError, TypeError):
        return None


def replace_placeholder_in_docx(doc, placeholder, value):
    clean_placeholder = "{" + placeholder + "}"
    all_placeholders = [
        clean_placeholder,
        "{**" + placeholder + "}",
        "{" + placeholder + "**}",
        "{**" + placeholder + "**}",
    ]

    for paragraph in doc.paragraphs:
        for ph in all_placeholders:
            if ph in paragraph.text:
                full_text = "".join(run.text for run in paragraph.runs)
                paragraph.clear()
                paragraph.add_run(full_text.replace(ph, str(value)))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for ph in all_placeholders:
                        if ph in paragraph.text:
                            full_text = "".join(run.text for run in paragraph.runs)
                            paragraph.clear()
                            paragraph.add_run(full_text.replace(ph, str(value)))


@document_bp.route("/api/analyze-template", methods=["POST"])
def analyze_template():
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No file selected"}), 400

        if not file.filename.endswith(".docx"):
            return jsonify({"error": "Only Word documents (.docx) are supported"}), 400

        doc = Document(BytesIO(file.read()))
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text + "\t"
                full_text += row_text.strip() + "\n"

        placeholders = re.findall(r"\{([^}]+)\}", full_text)
        cleaned_placeholders = []
        for ph in placeholders:
            clean = ph.replace("**", "").strip().replace(" ", "_")
            if clean and clean not in cleaned_placeholders:
                cleaned_placeholders.append(clean)

        return jsonify(
            {
                "placeholders": cleaned_placeholders,
                "count": len(cleaned_placeholders),
                "content": full_text,
            }
        )

    except Exception as e:
        print(f"Error analyzing template: {str(e)}", file=sys.stderr)
        print(traceback.format_exc(), file=sys.stderr)
        return jsonify({"error": f"Failed to analyze template: {str(e)}"}), 500


@document_bp.route("/api/generate-documents", methods=["POST"])
def generate_documents():
    task_id = None
    generated_docx_files = []  # Initialize the list

    try:
        if "template" not in request.files or "data" not in request.files:
            return jsonify({"error": "Template and data files are required"}), 400

        template_file = request.files["template"]
        data_file = request.files["data"]
        date_format = request.form.get("dateFormat", "DD MMMM YYYY")

        template_filename = template_file.filename
        template_name = os.path.splitext(template_filename)[0]

        df = pd.read_csv(data_file, dtype=str)
        df.columns = df.columns.str.strip().str.replace(" ", "_")
        records = df.to_dict("records")
        csv_headers = list(df.columns)

        # Extract all email addresses from all records
        all_emails = []
        email_columns = [col for col in df.columns if "email" in col.lower()]

        for record in records:
            for email_col in email_columns:
                email_value = record.get(email_col, "")
                if email_value and isinstance(email_value, str) and "@" in email_value:
                    # Handle multiple emails separated by commas/semicolons
                    emails = re.split(r"[;,\s]+", email_value.strip())
                    for email in emails:
                        if email and "@" in email:
                            all_emails.append(email.strip().lower())

        # Remove duplicates but preserve order
        seen = set()
        unique_emails = []
        for email in all_emails:
            if email not in seen:
                seen.add(email)
                unique_emails.append(email)

        task_id = str(uuid.uuid4())
        generation_progress[task_id] = {
            "total": len(records),
            "processed": 0,
            "current_file": "",
            "status": "processing",
        }

        temp_dir = tempfile.mkdtemp()
        word_dir = os.path.join(temp_dir, "word")
        pdf_dir = os.path.join(temp_dir, "pdf")
        os.makedirs(word_dir, exist_ok=True)
        os.makedirs(pdf_dir, exist_ok=True)

        template_content = template_file.read()

        # Store mapping of email to document filenames
        email_to_documents = {email: [] for email in unique_emails}

        # Step 1: Generate all DOCX
        generated_docx_files = []  # Initialize the list
        for i, record in enumerate(records):
            generation_progress[task_id]["processed"] = i + 1
            generation_progress[task_id][
                "current_file"
            ] = f"Processing record {i+1} of {len(records)}"

            doc = Document(BytesIO(template_content))
            record["Today"] = format_date(datetime.now(), date_format)
            
            formatted_record = {}
            for key, value in record.items():
                if value and is_date_field(key, value):
                    try:
                        date_obj = parse_date(value)
                        if date_obj:
                            formatted_record[key] = format_date(date_obj, date_format)
                        else:
                            formatted_record[key] = value
                    except:
                        formatted_record[key] = value
                else:
                    formatted_record[key] = value

            for key, value in formatted_record.items():
                replace_placeholder_in_docx(doc, key, value or "")

            applicant_name = record.get("Full_Name", f"applicant_{i+1}").replace(
                " ", "_"
            )
            docx_filename = f"{template_name}_{applicant_name}.docx"
            docx_path = os.path.join(word_dir, docx_filename)
            doc.save(docx_path)
            generated_docx_files.append(docx_path)

            # Map this document to all email addresses in this record
            for email_col in email_columns:
                email_value = record.get(email_col, "")
                if email_value and isinstance(email_value, str) and "@" in email_value:
                    emails = re.split(r"[;,\s]+", email_value.strip())
                    for email in emails:
                        if email and "@" in email:
                            clean_email = email.strip().lower()
                            if clean_email in email_to_documents:
                                email_to_documents[clean_email].append(docx_filename)

        # Update progress after DOCX generation
        generation_progress[task_id]["processed"] = len(records)
        generation_progress[task_id]["current_file"] = "Converting to PDF..."

        # Step 2: Convert DOCX → PDF in batch
        try:
            import pythoncom
            
            # Initialize COM only once
            pythoncom.CoInitialize()
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            for i, docx_path in enumerate(generated_docx_files):
                generation_progress[task_id]["processed"] = len(records) + i + 1
                generation_progress[task_id][
                    "current_file"
                ] = f"Converting to PDF {i+1} of {len(generated_docx_files)}"
                
                pdf_path = os.path.join(
                    pdf_dir, os.path.basename(docx_path).replace(".docx", ".pdf")
                )
                
                try:
                    doc = word.Documents.Open(docx_path)
                    doc.SaveAs(pdf_path, FileFormat=17)
                    doc.Close()
                except Exception as e:
                    print(f"Error converting {docx_path} to PDF: {str(e)}")
                    continue
            
            word.Quit()
            
        except Exception as e:
            generation_progress[task_id]["status"] = "error"
            generation_progress[task_id]["error"] = f"PDF conversion failed: {str(e)}"
            return jsonify({"error": str(e)}), 500
        finally:
            # Always uninitialize COM
            try:
                pythoncom.CoUninitialize()
            except:
                pass

        # Also create PDF filename mapping
        for email, doc_filenames in email_to_documents.items():
            pdf_filenames = [
                filename.replace(".docx", ".pdf") for filename in doc_filenames
            ]
            email_to_documents[email].extend(pdf_filenames)

        # Step 3: Zip both formats
        generation_progress[task_id]["current_file"] = "Creating zip files..."

        word_zip = os.path.join(temp_dir, "documents_word.zip")
        pdf_zip = os.path.join(temp_dir, "documents_pdf.zip")

        with zipfile.ZipFile(word_zip, "w") as zipf:
            for file_name in os.listdir(word_dir):
                zipf.write(os.path.join(word_dir, file_name), file_name)

        with zipfile.ZipFile(pdf_zip, "w") as zipf:
            for file_name in os.listdir(pdf_dir):
                zipf.write(os.path.join(pdf_dir, file_name), file_name)

        # Store all data
        generated_files[task_id] = {
            "word": word_zip,
            "pdf": pdf_zip,
            "headers": csv_headers,
            "emails": unique_emails,
            "records": records,
            "template_name": template_name,
            # New: mapping of email to document filenames
            "email_to_documents": email_to_documents,
            "extract_dir_word": word_dir,
            "extract_dir_pdf": pdf_dir,
        }

        generation_progress[task_id]["status"] = "completed"
        generation_progress[task_id]["processed"] = len(records) + len(
            generated_docx_files
        )
        generation_progress[task_id]["current_file"] = "All documents generated"

        return jsonify({"task_id": task_id}), 200

    except Exception as e:
            logging.error(f"Error in generate_documents: {str(e)}")
            logging.error(traceback.format_exc())
            
            if task_id and task_id in generation_progress:
                generation_progress[task_id]["status"] = "error"
                generation_progress[task_id]["error"] = str(e)
            
            return jsonify({"error": f"Failed to generate documents: {str(e)}"}), 500

# Other routes remain the same until email sending...


def ensure_outlook_initialized():
    """Ensure Outlook is properly initialized before sending emails"""
    try:
        import pythoncom

        pythoncom.CoInitialize()

        # Try to get the Outlook application instance
        try:
            outlook = win32com.client.GetActiveObject("Outlook.Application")
        except:
            # If not running, create a new instance
            outlook = win32com.client.Dispatch("Outlook.Application")

        # Force initialization by accessing a property
        _ = outlook.Version

        pythoncom.CoUninitialize()
        return True
    except Exception as e:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        print(f"Outlook initialization failed: {str(e)}")
        return False


def is_valid_email(email):
    """Validate email format"""
    import re
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def clean_email_address(email):
    """Clean email address by removing unwanted characters"""
    if not email:
        return ""
    
    # Remove any surrounding whitespace and convert to lowercase
    email = email.strip().lower()
    
    # Remove any non-ASCII characters that might cause issues
    email = email.encode('ascii', 'ignore').decode('ascii')
    
    # Remove any control characters
    email = ''.join(char for char in email if ord(char) >= 32)
    
    return email

def is_valid_email(email):
    """Validate email format"""
    import re
    if not email:
        return False
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def clean_email_address(email):
    """Clean email address by removing unwanted characters"""
    if not email:
        return ""
    
    # Remove any surrounding whitespace and convert to lowercase
    email = email.strip().lower()
    
    # Remove any non-ASCII characters that might cause issues
    email = email.encode('ascii', 'ignore').decode('ascii')
    
    # Remove any control characters
    email = ''.join(char for char in email if ord(char) >= 32)
    
    # Fix common typos
    email = email.replace('hotlook.com', 'hotmail.com')  # Common typo
    
    return email

def send_email_via_outlook(to_email, subject, body, attachment_paths=None, account=None, record=None, save_as_draft=False):
    """Send email via Outlook or save as draft if specified"""
    max_retries = 3
    base_delay = 2  # seconds
    
    # Clean and validate email first
    clean_to_email = clean_email_address(to_email)
    if not is_valid_email(clean_to_email):
        print(f"Invalid email format: {to_email} -> {clean_to_email}")
        return False
    
    for attempt in range(max_retries):
        try:
            import pythoncom
            pythoncom.CoInitialize()
            
            # Try to get the Outlook application instance
            try:
                outlook = win32com.client.GetActiveObject("Outlook.Application")
            except:
                # If not running, create a new instance
                outlook = win32com.client.Dispatch("Outlook.Application")
            
            # Create mail item
            mail = outlook.CreateItem(0)
            
            # Clean and prepare subject and body
            clean_subject = str(subject or "")
            clean_body = str(body or "")
            
            # Replace placeholders if record is provided
            if record:
                for key, value in record.items():
                    placeholder = "{" + key + "}"
                    safe_value = str(value or "")
                    clean_subject = clean_subject.replace(placeholder, safe_value)
                    clean_body = clean_body.replace(placeholder, safe_value)
            
            # Set basic email properties with cleaned email
            mail.To = clean_to_email
            mail.Subject = clean_subject
            
            # Set HTML body with proper encoding
            try:
                mail.HTMLBody = clean_body
            except Exception as html_error:
                # Fallback to plain text if HTML fails
                print(f"HTML body failed, using plain text: {html_error}")
                mail.Body = clean_body
            
            # Add attachments with validation
            if attachment_paths:
                for attachment_path in attachment_paths:
                    if os.path.exists(attachment_path):
                        try:
                            mail.Attachments.Add(attachment_path)
                        except Exception as attach_error:
                            print(f"Failed to add attachment {attachment_path}: {attach_error}")
                    else:
                        print(f"Attachment not found: {attachment_path}")
            
            # Try to use the specified account
            if account:
                try:
                    namespace = outlook.GetNamespace("MAPI")
                    for outlook_account in namespace.Accounts:
                        if outlook_account.SmtpAddress and outlook_account.SmtpAddress.lower() == account.lower():
                            try:
                                mail.SendUsingAccount = outlook_account
                            except AttributeError:
                                pass  # Ignore if not supported
                            break
                except Exception as account_error:
                    print(f"Warning: Could not set account {account}: {account_error}")
            
            # Save as draft or send
            if save_as_draft:
                mail.Save()
                print(f"Email saved as draft for: {clean_to_email}")
            else:
                mail.Send()
                print(f"Email sent to: {clean_to_email}")
            
            pythoncom.CoUninitialize()
            return True
            
        except Exception as e:
            error_msg = str(e)
            print(f"Attempt {attempt + 1} failed for {clean_to_email}: {error_msg}")
            
            try:
                pythoncom.CoUninitialize()
            except:
                pass
                
            if attempt < max_retries - 1:
                # Wait before retrying (exponential backoff)
                delay = base_delay * (2 ** attempt)
                print(f"Waiting {delay} seconds before retry...")
                time.sleep(delay)
            else:
                print(f"All attempts failed for {clean_to_email}")
                return False
    
    return False

@document_bp.route("/api/send-emails", methods=["POST"])
def send_emails():
    try:
        data = request.json
        document_task_id = data.get("task_id")
        email_subject = data.get("subject")
        email_body = data.get("body")
        email_format = data.get("format", "pdf")
        selected_account = data.get("account")
        send_multiple = data.get("send_multiple", False)
        save_drafts = data.get("save_drafts", True)  # Default to True

        if not document_task_id or document_task_id not in generated_files:
            return jsonify({"error": "Invalid document task ID"}), 400

        if not email_subject or not email_body:
            return jsonify({"error": "Email subject and body are required"}), 400

        # Get the generated files data
        file_data = generated_files[document_task_id]
        file_path = file_data[email_format]
        extract_dir = file_data[
            f"extract_dir_{'word' if email_format == 'word' else 'pdf'}"
        ]

        # Initialize email progress
        email_task_id = str(uuid.uuid4())
        email_progress[email_task_id] = {
            "total": len(file_data.get("emails", [])),
            "processed": 0,
            "success": 0,
            "failed": 0,
            "current_email": "",
            "status": "processing",
            "document_task_id": document_task_id,
        }

        # Start email sending in background thread
        thread = threading.Thread(
            target=send_emails_outlook_thread,
            args=(
                email_task_id,
                document_task_id,
                file_path,
                extract_dir,
                email_subject,
                email_body,
                email_format,
                selected_account,
                send_multiple,
                save_drafts,  # Add this parameter
            ),
        )
        thread.daemon = True
        thread.start()

        return jsonify({"task_id": email_task_id}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

def is_outlook_available():
    try:
        import pythoncom

        pythoncom.CoInitialize()

        outlook = win32com.client.Dispatch("Outlook.Application")
        namespace = outlook.GetNamespace("MAPI")
        # If we get here, Outlook is available

        pythoncom.CoUninitialize()
        return True
    except Exception as e:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        print(f"Outlook not available: {str(e)}")
        return False


# Update the send_emails_outlook_thread function to include retry logic
def send_emails_outlook_thread(task_id, document_task_id, file_path, extract_dir, 
                              subject, body, file_format, account=None, send_multiple=False, save_failed_as_draft=True):
    try:
        import pythoncom
        
        # Initialize COM for this thread
        pythoncom.CoInitialize()
        
        # Ensure Outlook is properly initialized before starting
        if not ensure_outlook_initialized():
            email_progress[task_id]["status"] = "error"
            email_progress[task_id]["error"] = "Microsoft Outlook is not available or not properly initialized."
            pythoncom.CoUninitialize()
            return
            
        email_data = generated_files[document_task_id]
        emails = email_data.get("emails", [])
        records = email_data.get("records", [])
        email_to_documents = email_data.get("email_to_documents", {})
        
        email_progress[task_id]["total"] = len(emails)
        
        # Initialize failure tracking
        email_failures[task_id] = []
        
        if len(emails) == 0:
            email_progress[task_id]["status"] = "completed"
            email_progress[task_id]["current_email"] = "No emails found to send"
            pythoncom.CoUninitialize()
            return
        
        # Create mapping for record data
        email_to_record = {}
        for i, record in enumerate(records):
            for key, value in record.items():
                if "email" in key.lower() and value:
                    email_values = re.split(r'[;,\s]+', str(value).strip())
                    for email in email_values:
                        if '@' in email:
                            clean_email = clean_email_address(email)
                            if clean_email:  # Only add valid emails
                                email_to_record[clean_email] = record
        
        # Process emails with rate limiting and better error handling
        processed_count = 0
        success_count = 0
        failure_count = 0
        
        for i, email_address in enumerate(emails):
            if email_progress[task_id]["status"] == "cancelled":
                break
                
            processed_count += 1
            email_progress[task_id]["processed"] = processed_count
            email_progress[task_id]["current_email"] = f"Preparing email {processed_count} of {len(emails)}: {email_address}"
            
            # Clean the email address
            clean_email = clean_email_address(email_address)
            if not is_valid_email(clean_email):
                failure_count += 1
                email_progress[task_id]["failed"] = failure_count
                email_failures[task_id].append({
                    "email": email_address,
                    "reason": f"Invalid email format: {email_address}",
                    "timestamp": datetime.now().isoformat(),
                    "attempts": 0
                })
                print(f"Skipping invalid email: {email_address}")
                continue
            
            # Get documents for this email
            document_filenames = email_to_documents.get(clean_email, [])
            
            # Prepare attachments
            if send_multiple:
                # Send ALL documents to this email
                all_documents = []
                for doc_email, docs in email_to_documents.items():
                    for doc in docs:
                        if doc.endswith(f'.{file_format}') and doc not in all_documents:
                            all_documents.append(doc)
                attachment_paths = [os.path.join(extract_dir, doc) for doc in all_documents 
                                 if os.path.exists(os.path.join(extract_dir, doc))]
            else:
                # Send only documents associated with this email
                relevant_docs = [doc for doc in document_filenames if doc.endswith(f'.{file_format}')]
                attachment_paths = [os.path.join(extract_dir, doc) for doc in relevant_docs 
                                 if os.path.exists(os.path.join(extract_dir, doc))]
            
            # Get record data for placeholder replacement
            record_data = email_to_record.get(clean_email, {"email": clean_email})
            
            # Update progress
            email_progress[task_id]["current_email"] = f"Sending to {clean_email} ({len(attachment_paths)} attachments)"
            
            # Try to send email with enhanced error handling
            max_email_retries = 2
            email_sent = False
            last_error = None
            
            for attempt in range(max_email_retries):
                try:
                    email_sent = send_email_via_outlook(
                        clean_email,
                        subject,
                        body,
                        attachment_paths,
                        account,
                        record_data
                    )
                    
                    if email_sent:
                        success_count += 1
                        email_progress[task_id]["success"] = success_count
                        email_progress[task_id]["current_email"] = f"✓ Sent to {clean_email}"
                        print(f"Successfully sent email to {clean_email}")
                        break
                    else:
                        last_error = "Unknown error - send function returned False"
                        time.sleep(2)  # Wait before retry
                        
                except Exception as e:
                    last_error = str(e)
                    print(f"Email attempt {attempt+1} failed for {clean_email}: {last_error}")
                    time.sleep(2)  # Wait before retry
            
            if not email_sent:
                # Save as draft if enabled
                if save_failed_as_draft:
                    try:
                        draft_saved = send_email_via_outlook(
                            clean_email,
                            subject,
                            body,
                            attachment_paths,
                            account,
                            record_data,
                            save_as_draft=True
                        )
                        if draft_saved:
                            last_error = f"Failed to send, but saved as draft in Outlook"
                            print(f"Failed email saved as draft for: {clean_email}")
                        else:
                            last_error = f"Failed to send and could not save as draft"
                    except Exception as draft_error:
                        last_error = f"Failed to send and draft saving failed: {str(draft_error)}"
                
                failure_count += 1
                email_progress[task_id]["failed"] = failure_count
                
                # Track the failure with more details
                email_failures[task_id].append({
                    "email": clean_email,
                    "reason": last_error or "Unknown error after all retries",
                    "timestamp": datetime.now().isoformat(),
                    "attempts": max_email_retries,
                    "attachments": len(attachment_paths),
                    "saved_as_draft": save_failed_as_draft and "draft" in last_error.lower()
                })
                print(f"Failed to send email to {clean_email} after {max_email_retries} attempts: {last_error}")
            
            # Rate limiting - wait between emails to avoid overwhelming Outlook
            if processed_count % EMAIL_RATE_LIMIT == 0 and processed_count < len(emails):
                email_progress[task_id]["current_email"] = f"Pausing for {EMAIL_TIME_PERIOD} seconds (rate limiting)..."
                time.sleep(EMAIL_TIME_PERIOD)
        
        # Final status update
        email_progress[task_id]["status"] = "completed"
        email_progress[task_id]["processed"] = len(emails)
        final_message = f"Completed. Success: {success_count}, Failed: {failure_count}"
        
        if failure_count > 0 and save_failed_as_draft:
            draft_count = sum(1 for f in email_failures[task_id] if f.get('saved_as_draft', False))
            final_message += f" - {draft_count} failed emails saved as drafts"
            
        email_progress[task_id]["current_email"] = final_message
        
        print(f"Email sending completed. Success: {success_count}, Failed: {failure_count}")
        
        pythoncom.CoUninitialize()
        
    except Exception as e:
        email_progress[task_id]["status"] = "error"
        email_progress[task_id]["error"] = f"Thread error: {str(e)}"
        print(f"Thread error: {str(e)}")
        traceback.print_exc()
        try:
            pythoncom.CoUninitialize()
        except:
            pass
# Add a new endpoint to download failed emails as CSV


@document_bp.route('/api/download-failed-emails/<task_id>')
def download_failed_emails(task_id):
    if task_id not in email_failures or not email_failures[task_id]:
        return jsonify({"error": "No failure data found for this task"}), 404
    
    # Create a CSV in memory using StringIO
    from io import StringIO
    
    si = StringIO()
    cw = csv.writer(si)
    
    # Write header
    cw.writerow(['Email', 'Reason', 'Timestamp', 'Attempts', 'Number of Attachments'])
    
    # Write data
    for failure in email_failures[task_id]:
        cw.writerow([
            failure['email'], 
            failure['reason'], 
            failure['timestamp'],
            failure.get('attempts', 'N/A'),
            failure.get('attachments', 'N/A')
        ])
    
    # Prepare response
    output = make_response(si.getvalue())
    output.headers["Content-Disposition"] = f"attachment; filename=failed_emails_{task_id}.csv"
    output.headers["Content-type"] = "text/csv"
    
    return output

@document_bp.route("/api/generation-progress/<task_id>")
def get_generation_progress(task_id):
    return jsonify(generation_progress.get(task_id, {}))


@document_bp.route("/api/download/<task_id>/<format>")
def download_file(task_id, format):
    if task_id not in generated_files:
        return jsonify({"error": "Task not found"}), 404

    if format not in ["word", "pdf"]:
        return jsonify({"error": "Invalid format"}), 400

    file_path = generated_files[task_id][format]
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404

    download_name = f"documents_{format}.zip"
    return send_file(file_path, as_attachment=True, download_name=download_name)


@document_bp.route("/api/cancel-generation", methods=["POST"])
def cancel_generation():
    task_id = request.json.get("task_id")
    if task_id in generation_progress:
        generation_progress[task_id]["status"] = "cancelled"
        return jsonify({"status": "cancelled"})
    return jsonify({"error": "Invalid task ID"}), 400


@document_bp.route("/api/get-placeholders/<task_id>")
def get_placeholders(task_id):
    if task_id not in generated_files:
        return jsonify({"error": "Invalid task ID"}), 400
    return jsonify({"placeholders": generated_files[task_id].get("headers", [])})


@document_bp.route("/api/get-outlook-accounts")
def get_outlook_accounts():
    try:
        import pythoncom

        pythoncom.CoInitialize()

        try:
            outlook = win32com.client.Dispatch("Outlook.Application")
            namespace = outlook.GetNamespace("MAPI")
            accounts = []

            for account in namespace.Accounts:
                try:
                    if hasattr(account, "SmtpAddress") and account.SmtpAddress:
                        accounts.append(account.SmtpAddress)
                except Exception as e:
                    print(f"Error reading account: {str(e)}")
                    continue

            pythoncom.CoUninitialize()

            if not accounts:
                return jsonify(
                    {
                        "accounts": [],
                        "warning": "Outlook is available but no accounts found",
                    }
                )

            return jsonify({"accounts": accounts})

        except Exception as e:
            pythoncom.CoUninitialize()
            error_msg = str(e)
            if "Invalid class string" in error_msg or "CLSID" in error_msg:
                return jsonify(
                    {"accounts": [], "error": "Microsoft Outlook is not installed"}
                )
            else:
                return jsonify(
                    {"accounts": [], "error": f"Outlook access error: {error_msg}"}
                )

    except Exception as e:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return jsonify(
            {"accounts": [], "error": f"COM initialization failed: {str(e)}"}
        )


@document_bp.route("/api/get-recipients/<task_id>")
def get_recipients(task_id):
    if task_id not in generated_files:
        return jsonify({"error": "Invalid task ID"}), 400
    return jsonify(
        {
            "emails": generated_files[task_id].get("emails", []),
            "count": len(generated_files[task_id].get("emails", [])),
        }
    )


@document_bp.route("/email-composer")
def serve_email_composer():
    return send_from_directory(".", "email_composer.html")


@document_bp.route("/api/email-progress/<task_id>")
def get_email_progress(task_id):
    progress = email_progress.get(task_id, {})

    if progress and "document_task_id" in progress:
        doc_task_id = progress["document_task_id"]
        if doc_task_id in generated_files:
            emails = generated_files[doc_task_id].get("emails", [])
            progress["total_emails_available"] = len(emails)

    return jsonify(progress)


@document_bp.route("/api/cancel-emails", methods=["POST"])
def cancel_emails():
    task_id = request.json.get("task_id")
    if task_id in email_progress:
        email_progress[task_id]["status"] = "cancelled"
        return jsonify({"status": "cancelled"})
    return jsonify({"error": "Invalid task ID"}), 400


# Add to app.py


@document_bp.route("/api/check-outlook-status")
def check_outlook_status():
    try:
        import pythoncom

        pythoncom.CoInitialize()

        # Try to create Outlook instance
        outlook = win32com.client.Dispatch("Outlook.Application")
        # Try to access a property to ensure it's working
        version = outlook.Version

        pythoncom.CoUninitialize()

        return jsonify({"available": True, "version": version})
    except Exception as e:
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return jsonify({"available": False, "error": str(e)})





